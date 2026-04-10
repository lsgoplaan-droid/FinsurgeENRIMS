from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('FinsurgeENRIMS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(75, 85, 180)

subtitle = doc.add_paragraph('Bank Demo — Q&A Guide & Talking Points')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(255, 102, 0)

date_para = doc.add_paragraph('2026-04-10')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(10)
date_para.runs[0].font.italic = True

doc.add_paragraph()  # Spacer

# Section 1: Critical Questions
doc.add_heading('Top 5 Critical Questions (Always Asked First)', level=1)
doc.add_paragraph('These are the questions that can make or break the deal. Be prepared with confident, detailed answers.')

# Q1
doc.add_heading('1. Data Security & Compliance: Where does data reside?', level=2)
q1 = doc.add_paragraph('Question: "Where is our customer and transaction data stored? Does it comply with RBI data localization requirements?"')
q1.runs[0].italic = True
doc.add_heading('Answer:', level=3)
points = [
    'All data resides exclusively in India (Azure Central India region)',
    'Complies with RBI Circular 2018 on data localization',
    'Encryption: PII fields (PAN, Aadhaar) encrypted at rest using AES-256',
    'API response masking: PAN shows as XXXXX1234, phone as +91XXXXX3210',
    '5-year audit trail maintained in tamper-proof, append-only format',
    'No data leaves India without explicit board approval'
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph('Follow-up if asked about backup/recovery:').runs[0].bold = True
backup_points = [
    'Daily automated backups with 30-day retention',
    'Point-in-time recovery capability (RTO < 4 hours)',
    'Backups stored in India region only',
    'Tested restore procedure validated quarterly'
]
for point in backup_points:
    doc.add_paragraph(point, style='List Bullet')

# Q2
doc.add_heading('2. Regulatory Filing: How does CTR/SAR auto-filing work?', level=2)
q2 = doc.add_paragraph('Question: "How do you ensure CTR/SAR accuracy? Can we review before filing? What\'s the audit trail?"')
q2.runs[0].italic = True
doc.add_heading('Answer:', level=3)
ctr_points = [
    'Workflow: True Positive case → Auto-generate CTR/SAR → Bank compliance officer reviews → Manual approval before filing → File to FIU-IND portal',
    'Pre-filled fields: Transaction details, customer KYC info, rule triggered, risk score, investigation findings',
    'Mandatory fields enforced: All required RBI/FIU fields validated before filing allowed',
    'Audit trail: Every CTR/SAR has complete history: created by (AI), reviewed by (compliance officer), filed by (authorized user), timestamp, IP address',
    'Backup copies: All filed CTR/SAR documents stored locally for 5 years',
    'Rejection handling: If FIU rejects, system auto-flags for re-investigation and resubmission'
]
for point in ctr_points:
    doc.add_paragraph(point, style='List Bullet')

# Q3
doc.add_heading('3. Alert Accuracy & False Positives: How do you prevent alert fatigue?', level=2)
q3 = doc.add_paragraph('Question: "With 26 rules, won\'t we be buried in alerts? What\'s your false positive rate?"')
q3.runs[0].italic = True
doc.add_heading('Answer:', level=3)
fp_points = [
    'Rules are tuned to your bank: Each detection rule has configurable thresholds',
    'Risk-based tuning: High-risk customers trigger stricter rules; low-risk only trigger high-confidence rules',
    'Demo baseline: 26 rules generate ~150 alerts/month on 10K customer base (~1.5% of transactions)',
    'Typical false positive rate: 30-40% (varies by rule and bank fraud patterns)',
    'Feedback loop: Bank marks false positives → system learns thresholds automatically',
    'Investigation Copilot: AI reduces investigation time to 30-45 min/alert (down from 4-6 hours manual)'
]
for point in fp_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph('Mitigation strategies:').runs[0].bold = True
mitig_points = [
    'Rules disabled by default; bank enables only high-confidence ones initially',
    'Velocity checks cross-reference with historical customer patterns',
    'Network analysis eliminates legitimate layering (e.g., payroll distribution)',
    'Merchant category whitelist prevents false alerts on legitimate suppliers'
]
for point in mitig_points:
    doc.add_paragraph(point, style='List Bullet')

# Q4
doc.add_heading('4. System Performance & SLA: What\'s your uptime guarantee?', level=2)
q4 = doc.add_paragraph('Question: "Can you handle our transaction volume? What happens if system goes down?"')
q4.runs[0].italic = True
doc.add_heading('Answer:', level=3)
perf_points = [
    'Current capacity: 10,000 TPS (transactions per second) per deployed instance',
    'Latency: Rule evaluation < 100ms, alert generation < 500ms',
    'Availability target: 99.5% uptime (4.3 hours downtime/month acceptable)',
    'Scale: Horizontal scaling via Kubernetes — add more instances for higher volume',
    'Peak hour handling: Auto-scaling triggers at 80% capacity; new instances online in < 2 minutes'
]
for point in perf_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph('Failover & redundancy:').runs[0].bold = True
failover_points = [
    'Primary-replica PostgreSQL with automatic failover (RTO 5-15 minutes)',
    'Redis cluster for session management (no session loss on node failure)',
    'Application load balancer across multiple zones (Azure Availability Zones)',
    'No single point of failure'
]
for point in failover_points:
    doc.add_paragraph(point, style='List Bullet')

# Q5
doc.add_heading('5. Integration & Onboarding: How long to go live?', level=2)
q5 = doc.add_paragraph('Question: "What\'s the implementation timeline? Do you integrate with our core banking system?"')
q5.runs[0].italic = True
doc.add_heading('Answer:', level=3)
onboard_points = [
    'Integration points: Your core banking system → REST API → FinsurgeENRIMS',
    'Data feed: Transaction files (CSV/JSON) or live API streaming',
    'Onboarding timeline: Week 1 (setup/training) → Week 2-3 (parallel run) → Week 4 (go-live)',
    'Data migration: Customer KYC profiles, historical transactions (12+ months recommended)',
    'No core banking changes required: System reads from your DB; no modifications to tables'
]
for point in onboard_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_page_break()

# Section 1b: Risk Score Calculation (TECHNICAL)
doc.add_heading('Technical Deep Dive: How Risk Scores Are Calculated', level=1)
doc.add_paragraph('This section explains the mathematical model behind the system\'s risk scoring.')

doc.add_heading('Transaction Risk Score (0–100)', level=2)
doc.add_paragraph('Calculated when a transaction is evaluated against detection rules:')

table_txn = doc.add_table(rows=6, cols=2)
table_txn.style = 'Light Grid Accent 1'
txn_rows = [
    ('Component', 'Contribution'),
    ('Base Score', '10 points (all transactions start here)'),
    ('Rule Severity Score', 'Critical: +40 | High: +25 | Medium: +15 | Low: +5'),
    ('Channel Risk', 'SWIFT: +10 | Branch/ATM: +5 | Internet/Mobile: +3 | POS: +2'),
    ('Customer Category Bonus', 'very_high: +35 | high: +20 | medium: +10 | low: 0'),
    ('Final Score', 'Sum of all factors, capped at 100'),
]
for i, (label, desc) in enumerate(txn_rows):
    cells = table_txn.rows[i].cells
    cells[0].text = label
    cells[1].text = desc

doc.add_paragraph('Example: Transaction from very_high risk customer, triggers high-severity rule, via SWIFT')
doc.add_paragraph('Score = 10 + 25 + 10 + 35 = 80/100', style='List Bullet')

doc.add_heading('Customer Risk Score (0–100)', level=2)
doc.add_paragraph('Updated dynamically whenever rules match a transaction. Previous score + adjustment (capped at 100).')

table_cust = doc.add_table(rows=5, cols=2)
table_cust.style = 'Light Grid Accent 1'
cust_rows = [
    ('Rule Severity', 'Risk Adjustment per Rule'),
    ('Critical', '+5 points'),
    ('High', '+3 points'),
    ('Medium', '+1.5 points'),
    ('Low', '+0.5 points'),
]
for i, (severity, adj) in enumerate(cust_rows):
    cells = table_cust.rows[i].cells
    cells[0].text = severity
    cells[1].text = adj

doc.add_paragraph('Example: Rajiv Saxena (current score 70) triggers critical rule twice')
doc.add_paragraph('New Score = min(70 + 5 + 5, 100) = 80/100', style='List Bullet')

doc.add_heading('Risk Category Auto-Assignment', level=2)
doc.add_paragraph('Categories assigned automatically based on customer risk score:')

table_cat = doc.add_table(rows=5, cols=2)
table_cat.style = 'Light Grid Accent 1'
cat_rows = [
    ('Score Range', 'Category'),
    ('75 – 100', 'very_high 🔴 (Highest monitoring)'),
    ('50 – 74', 'high 🟠 (Enhanced rules)'),
    ('25 – 49', 'medium 🟡 (Standard rules)'),
    ('0 – 24', 'low 🟢 (Minimal monitoring)'),
]
for i, (range_str, cat) in enumerate(cat_rows):
    cells = table_cat.rows[i].cells
    cells[0].text = range_str
    cells[1].text = cat

doc.add_paragraph('Key Insight: Customers with higher risk scores trigger stricter rules automatically, reducing false positives.', style='List Bullet')

doc.add_heading('Risk Appetite Metrics — CRO Dashboard', level=2)
doc.add_paragraph('The system monitors 5 portfolio-level metrics against CRO-defined thresholds. Status auto-updates in real-time:')

table_appetite = doc.add_table(rows=6, cols=4)
table_appetite.style = 'Light Grid Accent 1'
appetite_rows = [
    ('Metric', 'Calculation', 'Limit', 'Warning'),
    ('High-Risk Customer Exposure', '(very_high customers / total) × 100%', '15%', '12%'),
    ('Portfolio Average Risk Score', 'Average of all risk_scores (0–100)', '45', '35'),
    ('PEP Exposure', '(PEP customers / total) × 100%', '5%', '3%'),
    ('SLA Compliance Rate', '(on-time alerts / total) × 100%', '85%', '90%'),
    ('Flagged Transaction Volume (30d)', '(flagged amount / total) × 100%', '2%', '1.5%'),
]
for i, row_data in enumerate(appetite_rows):
    cells = table_appetite.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_heading('Status Indicators', level=3)
status_items = [
    '🟢 OK — Within warning threshold',
    '🟡 WARNING — Between warning and limit',
    '🔴 BREACH — Exceeds limit threshold',
]
for item in status_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('CRO can adjust all thresholds on-the-fly via dashboard without code changes or system restart.', style='List Bullet')

doc.add_page_break()

# Section 2: Key Advantage Table
doc.add_heading('Quick Reference — Your Key Advantages', level=1)

table = doc.add_table(rows=9, cols=2)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Question'
header_cells[1].text = 'Your Advantage'

# Data rows
advantages = [
    ('"Where is the data?"', 'India (Azure Central) — RBI compliant, encrypted at rest'),
    ('"CTR/SAR accuracy?"', 'Human approval workflow — compliance officer reviews before filing'),
    ('"False positives?"', 'Investigation Copilot AI reduces 4-6 hours to 30-45 minutes per alert'),
    ('"Uptime guarantee?"', '99.5% SLA, primary-replica failover, multi-zone redundancy'),
    ('"Go-live timeline?"', '4 weeks: setup (1w) + parallel run (2w) + cutover (1w)'),
    ('"vs. competitors?"', 'Only system with Copilot + Network analysis + RBI-native compliance'),
    ('"False positive rate?"', '30-40% (industry standard); Copilot reduces burden by 70%'),
    ('"ROI?"', 'Break-even in 6-12 months (fraud losses prevent implementation cost)'),
]

for i, (q, a) in enumerate(advantages, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = q
    row_cells[1].text = a

doc.add_page_break()

# Section 3: Demo Flow
doc.add_heading('Demo Flow Checklist', level=1)

checklist = [
    'Dashboard with metric drill-downs (30 alerts → drill to alert list)',
    'Alert investigation with Copilot AI suggestions',
    'Case creation from alert → assign to investigator',
    'Case disposition → CTR/SAR auto-generation',
    'Risk scoring explanation (why customer is high-risk)',
    'Network analysis for money mule detection (graph visualization)',
    'Full audit trail proof (RBI-ready, every action logged)'
]

for item in checklist:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Timing', level=2)
timing_points = [
    'Dashboard tour: 3-5 minutes',
    'Alert investigation: 3-5 minutes',
    'Case management: 3-5 minutes',
    'Risk scoring deep-dive: 5 minutes',
    'Network analysis: 3-5 minutes',
    'Q&A: 15-20 minutes',
    'Total: 30-45 minutes'
]
for item in timing_points:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Section 4: Closing Statements
doc.add_heading('Closing Statements', level=1)

doc.add_heading('If Positive Signals:', level=2)
positive = 'Based on your feedback, we\'re confident FinsurgeENRIMS will reduce your fraud losses, speed up investigations, and keep you compliant with RBI requirements. Next steps: 2-week pilot with your top 1000 customers, zero risk.'
doc.add_paragraph(positive, style='List Bullet')

doc.add_heading('If Skeptical:', level=2)
skeptical = 'I understand your concerns. Let\'s run a 2-week parallel test on a subset of your transactions, zero integration risk. If results match our claims, we move forward. If not, no obligation.'
doc.add_paragraph(skeptical, style='List Bullet')

doc.add_heading('Regulatory/Audit Concerns:', level=2)
regulatory = 'This audit trail is built specifically for RBI inspections. During any audit, you can export full transaction history with investigation notes, CTR/SAR proof, and SLA compliance metrics. You\'ll be RBI-ready on day one.'
doc.add_paragraph(regulatory, style='List Bullet')

doc.add_page_break()

# Section 5: Role-Based Talking Points
doc.add_heading('Role-Based Talking Points', level=1)

doc.add_heading('For CFO/Risk Committee:', level=2)
cfo_points = [
    'ROI in 6-12 months: fraud losses prevented offset implementation cost',
    'SLA compliance dashboard shows regulatory requirements met',
    'Audit trail proves ongoing monitoring (RBI audit defense)'
]
for point in cfo_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('For Compliance Officer:', level=2)
compliance_points = [
    'CTR/SAR auto-filing saves 20 hours/month manual work',
    'Risk scoring automatically prioritizes high-risk customers',
    '5-year audit trail maintained for every transaction'
]
for point in compliance_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('For IT Department:', level=2)
it_points = [
    'Cloud-native, scalable architecture (no infrastructure to manage)',
    'API-driven, integrates with existing systems',
    'Disaster recovery: RTO 4 hours, RPO 1 hour'
]
for point in it_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('For Operations/Analysts:', level=2)
ops_points = [
    'Investigation Copilot reduces investigation time from 4-6 hours to 30-45 minutes',
    'Network analysis shows money mule rings automatically',
    'Case dashboard and timeline make investigations 10x easier'
]
for point in ops_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_page_break()

# Footer
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run('Ready for Bank Demo | 2026-04-10')
footer_run.font.size = Pt(9)
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# Save
output_path = r"d:\Claude Projects\EnterpriseRiskSystem\FinsurgeENRIMS_BANK_DEMO_QA_UPDATED.docx"
doc.save(output_path)
print(f"Created: {output_path}")
